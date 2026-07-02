"""论文草稿 API 路由 —— 创建 / 编辑 / 生成大纲 / 生成章节 / 导出 docx"""
import io
import json
import logging
import os
import re
from datetime import datetime
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy.orm.attributes import flag_modified

from ..core.database import get_db
from ..models.draft import Draft
from ..models.outcome import Outcome
from ..models.project_design import ProjectDesign
from ..models.research_direction import ResearchDirection
from ..models.paper import Paper
from ..models.user import User
from ..schemas.draft import (
    DraftCreate, DraftUpdate, DraftOut, DraftOutline, PaperSection,
    GenerateOutlineRequest, GenerateChapterRequest,
    PAPER_CHAPTER_KEYS, PAPER_CHAPTER_LABELS,
    ChapterResult, AbstractResult, SuggestRefsResult, WritingPlanResult, WritingReviewResult, WritingRevisionResult,
    FullDraftReviewResult, FullDraftRevisionResult, FullDraftGenerateResult,
)
from ..schemas.compliance import ComplianceResult, ComplianceConfirmRequest
from ..agents.paper_writing_agent import paper_writing_agent
from ..agents.workflows import run_generate_chapter_workflow
from ..skills import SkillExecutionContext, get_default_skill_runtime
from ..tasks.paper_task import generate_chapter_task
from ..services.compliance_checker import check_draft
from ..services.evidence_retrieval_service import build_evidence_context, retrieve_project_evidence
from ..services.auth_dependency import get_current_user
from ..services.ownership import get_owned_draft, get_owned_project, query_owned_drafts
from ..core.config import settings

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/drafts", tags=["drafts"])

STORAGE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "storage", "drafts")


def _ensure_storage_dir():
    os.makedirs(STORAGE_DIR, exist_ok=True)


def _format_authors(authors_value) -> str:
    """将数据库中的作者字段格式化为可读文本。"""
    if not authors_value:
        return "佚名"
    if isinstance(authors_value, list):
        return ", ".join([str(a).strip() for a in authors_value[:3] if str(a).strip()]) or "佚名"
    if isinstance(authors_value, str):
        parts = [a.strip() for a in authors_value.split(";") if a.strip()]
        if not parts:
            parts = [a.strip() for a in authors_value.split(",") if a.strip()]
        return ", ".join(parts[:3]) if parts else "佚名"
    return str(authors_value)


def _format_jsonish_text(value) -> str:
    """把 JSON 字符串或列表整理成更适合给 LLM 的上下文文本。"""
    if value is None:
        return ""
    if isinstance(value, list):
        return "；".join(str(v).strip() for v in value if str(v).strip())
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return ""
        try:
            parsed = json.loads(raw)
        except Exception:
            return raw
        if isinstance(parsed, list):
            return "；".join(str(v).strip() for v in parsed if str(v).strip())
        if isinstance(parsed, dict):
            return json.dumps(parsed, ensure_ascii=False)
        return str(parsed)
    return str(value)


def _safe_filename(title: str, fallback: str, ext: str = "docx") -> str:
    """生成适合下载的文件名。"""
    cleaned = re.sub(r'[\\/:*?"<>|]+', "_", (title or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    if not cleaned:
        cleaned = fallback
    return f"{cleaned}.{ext}"


def _build_outcomes_summary(db: Session, project_id: UUID) -> str:
    """构建项目成果摘要文本"""
    outcomes = db.query(Outcome).filter(Outcome.project_id == project_id).all()
    if not outcomes:
        return "暂无上传成果 —— 论文中的实验和实现章节仅能编写设计方案和预期结果。"
    lines = [f"共 {len(outcomes)} 项成果："]
    for o in outcomes:
        lines.append(f"- [{o.outcome_type}] {o.name}: {o.description or '无描述'}")
    return "\n".join(lines)


def _build_project_context(db: Session, project_id: UUID) -> str:
    """构建项目背景上下文文本"""
    parts = []
    # 项目设计
    design = db.query(ProjectDesign).filter(ProjectDesign.project_id == project_id).order_by(
        ProjectDesign.created_at.desc()
    ).first()
    if design:
        topic = design.topic or "未指定"
        parts.append(f"研究题目：{topic}")
        if isinstance(design.content, dict):
            parts.append(json.dumps(design.content, ensure_ascii=False, indent=2))

    # 研究方向
    direction = db.query(ResearchDirection).filter(ResearchDirection.project_id == project_id).order_by(
        ResearchDirection.created_at.desc()
    ).first()
    if direction:
        parts.append(f"研究方向：{direction.title}")
        parts.append(f"研究问题：{_format_jsonish_text(direction.research_questions)}")
        parts.append(f"研究方法：{_format_jsonish_text(direction.methods)}")

    return "\n".join(parts) if parts else "暂无项目背景信息"


def _build_literature_context(
    db: Session,
    project_id: UUID,
    evidence_items: list[dict] | None = None,
) -> str:
    """构建文献上下文字符串"""
    papers = db.query(Paper).filter(Paper.project_id == project_id).limit(20).all()
    if evidence_items is None:
        evidence_items = retrieve_project_evidence(db, project_id, "", limit=12, min_confidence=70)
    if not papers and not evidence_items:
        return ""

    parts = []
    if papers:
        parts.append("已有文献：")
        for p in papers:
            authors = _format_authors(p.authors)
            parts.append(f"- {authors}. {p.title}. {p.venue or ''}, {p.year or ''}")

    evidence_context = build_evidence_context(evidence_items)
    if evidence_context:
        parts.append("")
        parts.append(evidence_context)
    return "\n".join(parts)


def _draft_to_out(draft: Draft) -> DraftOut:
    """Draft ORM → DraftOut 响应模型"""
    content = draft.content or {}
    sections = []
    for key in PAPER_CHAPTER_KEYS:
        ch = content.get(key, {})
        if isinstance(ch, dict):
            sections.append(PaperSection(
                key=key,
                title=ch.get("title", PAPER_CHAPTER_LABELS.get(key, key)),
                content=ch.get("content", ""),
                status=ch.get("status", "draft"),
            ))
        else:
            sections.append(PaperSection(
                key=key,
                title=PAPER_CHAPTER_LABELS.get(key, key),
                content="",
                status="draft",
            ))
    return DraftOut(
        id=str(draft.id),
        project_id=str(draft.project_id),
        title=draft.title,
        content=draft.content,
        references=draft.references,
        outline=draft.outline,
        version=draft.version or 1,
        sections=sections,
        created_at=draft.created_at,
        updated_at=draft.updated_at,
    )


def _build_full_draft_text(draft: Draft) -> str:
    """把现有章节内容拼成整篇正文，供整篇审查和修订使用。"""
    content = draft.content or {}
    parts = [f"# {draft.title}".strip()]
    for key in PAPER_CHAPTER_KEYS:
        chapter = content.get(key)
        if not isinstance(chapter, dict):
            continue
        chapter_text = str(chapter.get("content") or "").strip()
        if not chapter_text:
            continue
        title = chapter.get("title") or PAPER_CHAPTER_LABELS.get(key, key)
        parts.append(f"## {title}\n\n{chapter_text}")
    return "\n\n".join(parts).strip()


def _build_chapter_summaries(draft: Draft) -> list[dict]:
    """生成整篇审查所需的轻量章节摘要，不引入新表。"""
    content = draft.content or {}
    summaries = []
    for key in PAPER_CHAPTER_KEYS:
        chapter = content.get(key)
        title = PAPER_CHAPTER_LABELS.get(key, key)
        chapter_text = ""
        status = "draft"
        if isinstance(chapter, dict):
            title = chapter.get("title") or title
            chapter_text = str(chapter.get("content") or "")
            status = chapter.get("status") or status
        summaries.append({
            "key": key,
            "title": title,
            "length": len(chapter_text.strip()),
            "status": status,
        })
    return summaries


def _collect_draft_citations(draft: Draft) -> list[str]:
    """收集章节与 references 中已有的引用信息，供整篇审查判断证据状态。"""
    content = draft.content or {}
    citations: list[str] = []
    for key in PAPER_CHAPTER_KEYS:
        chapter = content.get(key)
        if isinstance(chapter, dict):
            citations.extend([str(item).strip() for item in chapter.get("citations", []) if str(item).strip()])
    for ref in draft.references or []:
        if isinstance(ref, dict):
            title = str(ref.get("title") or ref.get("citation_text") or "").strip()
            if title:
                citations.append(title)
        elif str(ref).strip():
            citations.append(str(ref).strip())
    return list(dict.fromkeys(citations))


def _apply_full_text_to_chapters(full_text: str, existing_content: dict | None) -> dict:
    """把保留了 Markdown 二级标题的全文回写到章节 JSON。"""
    content = dict(existing_content or {})
    title_to_key: dict[str, str] = {}
    for key in PAPER_CHAPTER_KEYS:
        labels = [PAPER_CHAPTER_LABELS.get(key, key)]
        existing = content.get(key)
        if isinstance(existing, dict) and existing.get("title"):
            labels.append(str(existing.get("title")))
        for label in labels:
            normalized = re.sub(r"\s+", "", str(label))
            if normalized:
                title_to_key[normalized] = key

    matches = list(re.finditer(r"^##\s+(.+?)\s*$", full_text or "", flags=re.MULTILINE))
    if not matches:
        return content

    for index, match in enumerate(matches):
        heading = match.group(1).strip()
        normalized_heading = re.sub(r"\s+", "", heading)
        chapter_key = title_to_key.get(normalized_heading)
        if not chapter_key:
            chapter_key = next(
                (
                    key
                    for label, key in title_to_key.items()
                    if normalized_heading in label or label in normalized_heading
                ),
                None,
            )
        if not chapter_key:
            continue
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(full_text)
        chapter_text = full_text[start:end].strip()
        current = content.get(chapter_key) if isinstance(content.get(chapter_key), dict) else {}
        content[chapter_key] = {
            **current,
            "title": current.get("title") or PAPER_CHAPTER_LABELS.get(chapter_key, heading),
            "content": chapter_text,
            "status": "edited",
        }

    return content
# ---- CRUD ----

@router.post("/", response_model=DraftOut, status_code=201)
def create_draft(
    payload: DraftCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """创建论文草稿"""
    try:
        project = get_owned_project(payload.project_id, current_user, db)
        draft = Draft(
            project_id=project.id,
            title=payload.title,
            content={},
            references=[],
            version=1,
        )
        db.add(draft)
        db.commit()
        db.refresh(draft)
        return _draft_to_out(draft)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"创建草稿失败: {str(e)}")


@router.get("/", response_model=list[DraftOut])
def list_drafts(
    project_id: str | None = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """列出论文草稿，可按项目过滤"""
    q = query_owned_drafts(db, current_user)
    if project_id:
        project = get_owned_project(project_id, current_user, db)
        q = q.filter(Draft.project_id == project.id)
    drafts = q.order_by(Draft.updated_at.desc()).all()
    return [_draft_to_out(d) for d in drafts]


@router.get("/{draft_id}", response_model=DraftOut)
def get_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取草稿完整内容"""
    draft = get_owned_draft(draft_id, current_user, db)
    return _draft_to_out(draft)


@router.patch("/{draft_id}", response_model=DraftOut)
def update_draft(
    draft_id: UUID,
    payload: DraftUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """更新草稿（保存用户编辑）"""
    draft = get_owned_draft(draft_id, current_user, db)
    try:
        for field, value in payload.model_dump(exclude_unset=True).items():
            if value is not None:
                setattr(draft, field, value)
        db.commit()
        db.refresh(draft)
        return _draft_to_out(draft)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"更新草稿失败: {str(e)}")


@router.delete("/{draft_id}", status_code=204)
def delete_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """删除草稿"""
    draft = get_owned_draft(draft_id, current_user, db)
    try:
        db.delete(draft)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除草稿失败: {str(e)}")


# ---- AI 生成 ----

@router.post("/{draft_id}/outline", response_model=DraftOutline)
def generate_outline(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """生成论文大纲"""
    draft = get_owned_draft(draft_id, current_user, db)

    project_context = _build_project_context(db, draft.project_id)
    outcomes_summary = _build_outcomes_summary(db, draft.project_id)
    literature_context = _build_literature_context(db, draft.project_id)

    skill_runtime = get_default_skill_runtime()
    skill_definition = skill_runtime.router.resolve(domain="paper", action="generate_outline")
    result = skill_runtime.executor.execute(
        skill_definition.id,
        {
            "project_context": project_context,
            "outcomes_summary": outcomes_summary,
            "literature_context": literature_context,
        },
        context=SkillExecutionContext(
            user_id=str(current_user.id),
            project_id=str(draft.project_id),
            draft_id=str(draft.id),
            state={"writing_agent": paper_writing_agent},
        ),
    )
    outline_result = result.output

    # 保存大纲到草稿
    draft.outline = outline_result
    draft.title = outline_result.get("suggested_title", draft.title)
    db.commit()

    from ..schemas.draft import ChapterOutline
    chapters = [
        ChapterOutline(
            key=ch.get("key", ""),
            title=ch.get("title", ""),
            subsections=ch.get("subsections", []),
        )
        for ch in outline_result.get("chapters", [])
    ]
    return DraftOutline(
        chapters=chapters,
        suggested_title=outline_result.get("suggested_title"),
        notes=outline_result.get("notes"),
    )


@router.post("/{draft_id}/plan", response_model=WritingPlanResult)
def generate_writing_plan(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """生成论文写作计划。"""
    draft = get_owned_draft(draft_id, current_user, db)

    project_context = _build_project_context(db, draft.project_id)
    outcomes_summary = _build_outcomes_summary(db, draft.project_id)
    literature_context = _build_literature_context(db, draft.project_id)

    skill_runtime = get_default_skill_runtime()
    skill_definition = skill_runtime.router.resolve(domain="paper", action="plan")
    result = skill_runtime.executor.execute(
        skill_definition.id,
        {
            "project_context": project_context,
            "outcomes_summary": outcomes_summary,
            "literature_context": literature_context,
        },
        context=SkillExecutionContext(
            user_id=str(current_user.id),
            project_id=str(draft.project_id),
            draft_id=str(draft.id),
            state={"writing_agent": paper_writing_agent},
        ),
    )
    plan_result = result.output

    content = dict(draft.content or {})
    content["_writing_plan"] = plan_result
    draft.content = content
    flag_modified(draft, "content")
    db.commit()

    return WritingPlanResult(**plan_result)


@router.post("/{draft_id}/full-generate", response_model=DraftOut)
def generate_full_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """生成整篇初稿，保留已手工编辑章节。"""
    draft = get_owned_draft(draft_id, current_user, db)

    project_context = _build_project_context(db, draft.project_id)
    outcomes_summary = _build_outcomes_summary(db, draft.project_id)
    literature_context = _build_literature_context(db, draft.project_id)

    result = paper_writing_agent.generate_full_draft(
        project_context=project_context,
        outcomes_summary=outcomes_summary,
        literature_context=literature_context,
        existing_outline=draft.outline or None,
        existing_chapters=draft.content or {},
    )

    draft.title = result.get("suggested_title", draft.title)
    draft.outline = result.get("outline", draft.outline or {})
    draft.content = result.get("content", draft.content or {})
    draft.version = (draft.version or 1) + 1
    flag_modified(draft, "content")
    db.commit()
    db.refresh(draft)

    return _draft_to_out(draft)


@router.post("/{draft_id}/chapters/{chapter_key}/review", response_model=WritingReviewResult)
def review_chapter(
    draft_id: UUID,
    chapter_key: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """审查当前章节的结构、依据与表达问题。"""
    if chapter_key not in PAPER_CHAPTER_KEYS:
        raise HTTPException(status_code=400, detail=f"无效的章节标识: {chapter_key}")

    draft = get_owned_draft(draft_id, current_user, db)
    chapter_data = (draft.content or {}).get(chapter_key)
    if not isinstance(chapter_data, dict) or not str(chapter_data.get("content") or "").strip():
        raise HTTPException(status_code=400, detail="当前章节暂无内容，无法审查")

    evidence_context = _build_literature_context(db, draft.project_id)
    skill_runtime = get_default_skill_runtime()
    skill_definition = skill_runtime.router.resolve(domain="paper", action="review_chapter")
    result = skill_runtime.executor.execute(
        skill_definition.id,
        {
            "chapter_key": chapter_key,
            "chapter_title": chapter_data.get("title", PAPER_CHAPTER_LABELS.get(chapter_key, chapter_key)),
            "chapter_content": chapter_data.get("content", ""),
            "citations": chapter_data.get("citations", []),
            "evidence_context": evidence_context,
        },
        context=SkillExecutionContext(
            user_id=str(current_user.id),
            project_id=str(draft.project_id),
            draft_id=str(draft.id),
            state={"writing_agent": paper_writing_agent},
        ),
    )
    review_result = result.output

    content = dict(draft.content or {})
    content["_chapter_reviews"] = {
        **(content.get("_chapter_reviews") if isinstance(content.get("_chapter_reviews"), dict) else {}),
        chapter_key: review_result,
    }
    draft.content = content
    flag_modified(draft, "content")
    db.commit()

    return WritingReviewResult(**review_result)


@router.post("/{draft_id}/chapters/{chapter_key}/revise", response_model=WritingRevisionResult)
def revise_chapter(
    draft_id: UUID,
    chapter_key: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """根据当前章节审查结果执行定向修订。"""
    if chapter_key not in PAPER_CHAPTER_KEYS:
        raise HTTPException(status_code=400, detail=f"无效的章节标识: {chapter_key}")

    draft = get_owned_draft(draft_id, current_user, db)
    chapter_data = (draft.content or {}).get(chapter_key)
    if not isinstance(chapter_data, dict) or not str(chapter_data.get("content") or "").strip():
        raise HTTPException(status_code=400, detail="当前章节暂无内容，无法修订")

    review_map = (draft.content or {}).get("_chapter_reviews")
    review_data = review_map.get(chapter_key) if isinstance(review_map, dict) else None
    if not isinstance(review_data, dict):
        raise HTTPException(status_code=400, detail="请先执行章节审查，再进行定向修订")

    evidence_context = _build_literature_context(db, draft.project_id)
    skill_runtime = get_default_skill_runtime()
    skill_definition = skill_runtime.router.resolve(domain="paper", action="apply_revision")
    result = skill_runtime.executor.execute(
        skill_definition.id,
        {
            "chapter_key": chapter_key,
            "chapter_title": chapter_data.get("title", PAPER_CHAPTER_LABELS.get(chapter_key, chapter_key)),
            "chapter_content": chapter_data.get("content", ""),
            "issues": review_data.get("issues", []),
            "focus_areas": review_data.get("focus_areas", []),
            "citations": chapter_data.get("citations", []),
            "evidence_context": evidence_context,
        },
        context=SkillExecutionContext(
            user_id=str(current_user.id),
            project_id=str(draft.project_id),
            draft_id=str(draft.id),
            state={"writing_agent": paper_writing_agent},
        ),
    )
    revision_result = result.output

    content = dict(draft.content or {})
    content[chapter_key] = {
        "title": revision_result.get("title", chapter_data.get("title", PAPER_CHAPTER_LABELS.get(chapter_key, chapter_key))),
        "content": revision_result.get("content", chapter_data.get("content", "")),
        "status": "edited",
        "citations": revision_result.get("citations", chapter_data.get("citations", [])),
        "data_based": revision_result.get("data_based", chapter_data.get("data_based", False)),
    }
    content["_chapter_revisions"] = {
        **(content.get("_chapter_revisions") if isinstance(content.get("_chapter_revisions"), dict) else {}),
        chapter_key: revision_result,
    }
    draft.content = content
    draft.version = (draft.version or 1) + 1
    flag_modified(draft, "content")
    db.commit()

    return WritingRevisionResult(**revision_result)


@router.post("/{draft_id}/review-full", response_model=FullDraftReviewResult)
def review_full_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """审查整篇论文的结构、证据、衔接和重复风险。"""
    draft = get_owned_draft(draft_id, current_user, db)
    full_text = _build_full_draft_text(draft)
    if not full_text or "##" not in full_text:
        raise HTTPException(status_code=400, detail="当前草稿暂无可审查的整篇正文")

    evidence_context = _build_literature_context(db, draft.project_id)
    skill_runtime = get_default_skill_runtime()
    skill_definition = skill_runtime.router.resolve(domain="paper", action="review_full")
    result = skill_runtime.executor.execute(
        skill_definition.id,
        {
            "draft_title": draft.title,
            "full_text": full_text,
            "chapter_summaries": _build_chapter_summaries(draft),
            "citations": _collect_draft_citations(draft),
            "evidence_context": evidence_context,
        },
        context=SkillExecutionContext(
            user_id=str(current_user.id),
            project_id=str(draft.project_id),
            draft_id=str(draft.id),
            state={"writing_agent": paper_writing_agent},
        ),
    )
    review_result = result.output

    content = dict(draft.content or {})
    content["_full_review"] = review_result
    draft.content = content
    flag_modified(draft, "content")
    db.commit()

    return FullDraftReviewResult(**review_result)


@router.post("/{draft_id}/revise-full", response_model=FullDraftRevisionResult)
def revise_full_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """根据整篇审查结果执行整篇轻量修订并回写章节正文。"""
    draft = get_owned_draft(draft_id, current_user, db)
    content = dict(draft.content or {})
    review_data = content.get("_full_review")
    if not isinstance(review_data, dict):
        raise HTTPException(status_code=400, detail="请先执行整篇审查，再进行整篇修订")

    full_text = _build_full_draft_text(draft)
    if not full_text or "##" not in full_text:
        raise HTTPException(status_code=400, detail="当前草稿暂无可修订的整篇正文")

    evidence_context = _build_literature_context(db, draft.project_id)
    skill_runtime = get_default_skill_runtime()
    skill_definition = skill_runtime.router.resolve(domain="paper", action="revise_full")
    result = skill_runtime.executor.execute(
        skill_definition.id,
        {
            "draft_title": draft.title,
            "full_text": full_text,
            "issues": review_data.get("issues", []),
            "focus_areas": review_data.get("focus_areas", []),
            "citations": _collect_draft_citations(draft),
            "evidence_context": evidence_context,
        },
        context=SkillExecutionContext(
            user_id=str(current_user.id),
            project_id=str(draft.project_id),
            draft_id=str(draft.id),
            state={"writing_agent": paper_writing_agent},
        ),
    )
    revision_result = result.output

    updated_content = _apply_full_text_to_chapters(revision_result.get("full_text", ""), content)
    updated_content["_full_review"] = review_data
    updated_content["_full_revision"] = revision_result
    draft.content = updated_content
    draft.title = revision_result.get("title") or draft.title
    draft.version = (draft.version or 1) + 1
    flag_modified(draft, "content")
    db.commit()

    return FullDraftRevisionResult(**revision_result)


@router.post("/{draft_id}/chapters/{chapter_key}", response_model=ChapterResult)
def generate_chapter(
    draft_id: UUID,
    chapter_key: str,
    payload: GenerateChapterRequest | None = None,
    async_mode: bool = Query(False, alias="async", description="是否异步生成"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """生成单个章节内容"""
    if chapter_key not in PAPER_CHAPTER_KEYS:
        raise HTTPException(status_code=400, detail=f"无效的章节标识: {chapter_key}")

    draft = get_owned_draft(draft_id, current_user, db)

    # 异步模式
    if async_mode:
        task = generate_chapter_task.delay(
            draft_id=str(draft.id),
            chapter_key=chapter_key,
            literature_context=_build_literature_context(db, draft.project_id),
        )
        return {"task_id": task.id, "status": "pending", "chapter_key": chapter_key}

    try:
        result = run_generate_chapter_workflow(
            db=db,
            draft=draft,
            chapter_key=chapter_key,
            user_id=str(current_user.id),
            writing_agent=paper_writing_agent,
            record_db=db,
        )
    except ValueError as e:
        raise HTTPException(status_code=422, detail=f"章节生成结果缺少可验证依据: {e}")

    return ChapterResult(
        chapter_key=chapter_key,
        title=result.get("title", PAPER_CHAPTER_LABELS.get(chapter_key, chapter_key)),
        content=result.get("content", ""),
        status="generated",
        citations=result.get("citations", []),
        data_based=result.get("data_based", False),
    )


@router.post("/{draft_id}/abstract", response_model=AbstractResult)
def generate_abstract(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """根据全文生成中英文摘要"""
    draft = get_owned_draft(draft_id, current_user, db)

    result = paper_writing_agent.generate_abstract(draft.content or {})
    return AbstractResult(
        abstract_cn=result.get("abstract_cn", ""),
        abstract_en=result.get("abstract_en", ""),
        keywords_cn=result.get("keywords_cn", []),
        keywords_en=result.get("keywords_en", []),
    )


@router.post("/{draft_id}/suggest-refs", response_model=SuggestRefsResult)
def suggest_references(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """推荐补充参考文献"""
    draft = get_owned_draft(draft_id, current_user, db)

    papers = db.query(Paper).filter(Paper.project_id == draft.project_id).all()
    existing_lit = [
        {
            "title": p.title,
            "authors": p.authors,
            "year": p.year,
            "venue": p.venue,
            "doi": p.doi,
        }
        for p in papers
    ]

    result = paper_writing_agent.suggest_references(draft.content or {}, existing_lit)
    return SuggestRefsResult(
        suggested_references=result.get("suggested_references", []),
        notes=result.get("notes"),
    )


# ---- 学术合规检查 ----

@router.post("/{draft_id}/check-compliance", response_model=ComplianceResult)
def check_compliance(
    draft_id: UUID,
    enable_ai: bool = Query(False, description="是否启用 AI 深度语义检查"),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """对论文草稿运行合规检查（规则检查始终执行，AI 检查可选）。"""
    draft = get_owned_draft(draft_id, current_user, db)

    outcomes = db.query(Outcome).filter(Outcome.project_id == draft.project_id).all()
    papers = db.query(Paper).filter(Paper.project_id == draft.project_id).all()

    result = check_draft(
        draft=draft,
        outcomes=outcomes,
        papers=papers,
        enable_ai=enable_ai,
        api_key=settings.DEEPSEEK_API_KEY if enable_ai else None,
        base_url=settings.DEEPSEEK_BASE_URL if enable_ai else None,
        model=settings.DEEPSEEK_MODEL if enable_ai else None,
    )

    # 持久化合规结果
    content = dict(draft.content or {})
    content["_compliance"] = result.model_dump(mode="json")
    draft.content = content
    flag_modified(draft, "content")
    db.commit()

    return result


@router.post("/{draft_id}/confirm-compliance")
def confirm_compliance(
    draft_id: UUID,
    payload: ComplianceConfirmRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """用户确认/忽略/修正某个合规 issue。"""
    draft = get_owned_draft(draft_id, current_user, db)

    content = dict(draft.content or {})
    compliance = content.get("_compliance", {})
    chapters = compliance.get("chapters", {})

    chapter = chapters.get(payload.chapter_key)
    if not chapter:
        raise HTTPException(status_code=404, detail=f"章节 {payload.chapter_key} 无检查结果")

    issues = chapter.get("issues", [])
    if payload.issue_index < 0 or payload.issue_index >= len(issues):
        raise HTTPException(status_code=400, detail=f"issue_index 超出范围 (0-{len(issues) - 1})")

    issues[payload.issue_index]["user_action"] = payload.action
    issues[payload.issue_index]["confirmed_at"] = datetime.utcnow().isoformat()

    # 检查该章节是否所有 error 都已确认
    has_unconfirmed_error = any(
        i.get("severity") == "error" and i.get("user_action") not in ("accept",)
        for i in issues
    )
    chapter["confirmed"] = not has_unconfirmed_error
    if chapter["confirmed"]:
        chapter["confirmed_at"] = datetime.utcnow().isoformat()

    content["_compliance"] = compliance
    draft.content = content
    flag_modified(draft, "content")
    db.commit()

    return {"status": "ok", "chapter_key": payload.chapter_key, "chapter_confirmed": chapter["confirmed"]}


@router.get("/{draft_id}/compliance-status")
def get_compliance_status(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """获取论文草稿的合规检查状态。"""
    draft = get_owned_draft(draft_id, current_user, db)

    content = draft.content or {}
    compliance = content.get("_compliance")
    if not compliance:
        return {"checked": False, "message": "尚未执行合规检查"}

    return compliance


# ---- 导出 ----

@router.get("/{draft_id}/preview")
def preview_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """拼接返回完整论文文本预览"""
    draft = get_owned_draft(draft_id, current_user, db)

    content = draft.content or {}
    parts = [f"# {draft.title}\n\n"]
    for key in PAPER_CHAPTER_KEYS:
        ch = content.get(key, {})
        if isinstance(ch, dict) and ch.get("content"):
            parts.append(f"## {PAPER_CHAPTER_LABELS.get(key, key)}\n\n")
            parts.append(ch["content"])
            parts.append("\n\n---\n\n")

    return {"title": draft.title, "full_text": "".join(parts), "version": draft.version}


@router.get("/{draft_id}/download")
def download_draft(
    draft_id: UUID,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    format: str = Query("docx", pattern="^(docx|pdf)$"),
):
    """导出论文草稿为 docx 或 pdf 文件（MinIO 优先，本地 fallback）"""
    draft = get_owned_draft(draft_id, current_user, db)

    if format == "pdf":
        content_type = "application/pdf"
        ext = "pdf"
        buf = io.BytesIO(_build_pdf(draft))
    else:
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
        buf = _build_docx(draft)

    filename = _safe_filename(draft.title, f"draft_{draft_id}", ext)
    object_key = f"drafts/draft_{draft_id}.{ext}"
    file_bytes = buf.read()

    # MinIO 优先
    try:
        from ..services.upload_service import get_object_stream as gos
        minio_result = gos(object_key)
        if minio_result is not None:
            stream, size, ct = minio_result
            return StreamingResponse(
                stream,
                media_type=content_type,
                headers={
                    "Content-Disposition": f'attachment; filename="{filename}"',
                    "Content-Length": str(size),
                },
            )
    except Exception:
        pass

    # 生成并保存到本地
    _ensure_storage_dir()
    filepath = os.path.join(STORAGE_DIR, f"draft_{draft_id}.{ext}")
    with open(filepath, "wb") as f:
        f.write(file_bytes)

    # 异步保存到 MinIO
    try:
        from ..services.upload_service import save_bytes as sb
        sb(file_bytes, object_key, content_type=content_type)
    except Exception:
        pass

    return FileResponse(
        filepath,
        media_type=content_type,
        filename=filename,
    )


# ---- PDF 渲染 ----

def _build_pdf(draft: Draft) -> bytes:
    """将论文草稿渲染为 PDF 字节"""
    from ..services.pdf_builder import PdfBuilder

    pdf = PdfBuilder(draft.title)
    content = draft.content or {}
    for key in PAPER_CHAPTER_KEYS:
        ch = content.get(key, {})
        ch_title = PAPER_CHAPTER_LABELS.get(key, key)
        ch_content = ch.get("content", "") if isinstance(ch, dict) else str(ch)

        pdf.add_heading(ch_title, level=1)
        if ch_content:
            for para_text in ch_content.strip().split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                pdf.add_body(para_text)

    return pdf.output()


# ---- DOCX 渲染 ----

def _build_docx(draft: Draft) -> io.BytesIO:
    """将论文草稿渲染为 docx 文件"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_para.paragraph_format.space_after = Pt(30)
    run = title_para.add_run(draft.title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 章节
    content = draft.content or {}
    for key in PAPER_CHAPTER_KEYS:
        ch = content.get(key, {})
        ch_title = PAPER_CHAPTER_LABELS.get(key, key)
        ch_content = ch.get("content", "") if isinstance(ch, dict) else str(ch)

        heading = doc.add_heading(ch_title, level=1)
        for run in heading.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(16)

        if ch_content:
            for para_text in ch_content.strip().split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(para_text)
                run.font.name = "宋体"
                run.font.size = Pt(12)
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
