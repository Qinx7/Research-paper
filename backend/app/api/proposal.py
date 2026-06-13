"""开题报告 API 路由 —— 生成 / 查看 / 下载 docx"""
import io
import logging
import os
import re

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session

from ..core.database import get_db
from ..models.proposal import Proposal
from ..models.project_design import ProjectDesign
from ..models.research_direction import ResearchDirection
from ..schemas.proposal import (
    ProposalGenerateRequest,
    ProposalOut,
    ProposalSection,
    SECTION_KEYS,
    SECTION_LABELS,
)
from ..agents.proposal_agent import proposal_agent
from ..tasks.proposal_task import generate_proposal_task
from ..services.grounding_guard import collect_allowed_references_from_design

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/proposal", tags=["proposal"])

STORAGE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "storage", "proposals")


def _ensure_storage_dir():
    os.makedirs(STORAGE_DIR, exist_ok=True)


def _safe_filename(title: str, fallback: str, ext: str = "docx") -> str:
    """生成适合下载的文件名。"""
    cleaned = re.sub(r'[\\/:*?"<>|]+', "_", (title or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    if not cleaned:
        cleaned = fallback
    return f"{cleaned}.{ext}"


def _build_pdf(proposal: Proposal) -> bytes:
    """将开题报告渲染为 PDF 字节"""
    from ..services.pdf_builder import PdfBuilder

    pdf = PdfBuilder(proposal.title)
    sections_data = proposal.content or {}
    for key in SECTION_KEYS:
        sec = sections_data.get(key, {})
        sec_title = SECTION_LABELS.get(key, key)
        sec_content = sec.get("content", "") if isinstance(sec, dict) else str(sec)

        pdf.add_heading(sec_title, level=1)
        if sec_content:
            for para_text in sec_content.strip().split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                pdf.add_body(para_text)

    return pdf.output()


def _build_docx(proposal: Proposal) -> io.BytesIO:
    """将开题报告内容渲染为 docx 文件，返回 BytesIO 流"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # ---- 封面标题 ----
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_para.paragraph_format.space_after = Pt(30)
    run = title_para.add_run(proposal.title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(40)
    run = sub_para.add_run("开题报告")
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 分隔线
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run("─" * 40)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    doc.add_paragraph()  # 空行

    # ---- 各章节正文 ----
    sections_data = proposal.content or {}
    for key in SECTION_KEYS:
        sec = sections_data.get(key, {})
        sec_title = SECTION_LABELS.get(key, key)
        sec_content = sec.get("content", "") if isinstance(sec, dict) else str(sec)

        # 章节标题
        heading = doc.add_heading(sec_title, level=1)
        for run in heading.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)

        # 章节内容（按段落拆分）
        paragraphs = sec_content.strip().split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(para_text)
            run.font.name = "宋体"
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _proposal_to_out(proposal: Proposal) -> ProposalOut:
    """将 ORM 对象转为响应模型"""
    content = proposal.content or {}
    sections = []
    for key in SECTION_KEYS:
        sec = content.get(key, {})
        sec_title = SECTION_LABELS.get(key, key)
        sec_content = sec.get("content", "") if isinstance(sec, dict) else str(sec)
        sections.append(ProposalSection(key=key, title=sec_title, content=sec_content))
    return ProposalOut(
        id=str(proposal.id),
        project_id=str(proposal.project_id) if proposal.project_id else None,
        design_id=str(proposal.design_id) if proposal.design_id else None,
        title=proposal.title,
        sections=sections,
        docx_path=proposal.docx_path,
        created_at=proposal.created_at,
    )


@router.post("/generate")
def generate_proposal(
    payload: ProposalGenerateRequest,
    db: Session = Depends(get_db),
    async_mode: bool = Query(False, alias="async", description="是否异步生成（Celery 任务）"),
):
    """根据项目设计生成开题报告（12 章节完整版）。

    设置 ?async=true 则立即返回 task_id，客户端可轮询 GET /api/tasks/{task_id} 获取结果。
    """
    # --- 异步模式：分发到 Celery ---
    if async_mode:
        # 加载项目设计（同步部分，必须在这里完成）
        design = db.query(ProjectDesign).filter(ProjectDesign.id == payload.design_id).first()
        if not design:
            raise HTTPException(status_code=404, detail="项目设计不存在")

        direction = None
        if design.direction_id:
            rd = db.query(ResearchDirection).filter(ResearchDirection.id == design.direction_id).first()
            if rd:
                direction = {
                    "title": rd.title,
                    "background": rd.background,
                    "research_questions": rd.research_questions,
                    "methods": rd.methods,
                    "expected_outputs": rd.expected_outputs,
                    "innovation": rd.innovation,
                    "content": rd.content,
                }

        literature_context = ""
        if isinstance(design.content, dict):
            lit_review = design.content.get("literature_review", {})
            if isinstance(lit_review, dict):
                parts = []
                refs = lit_review.get("key_references", [])
                if refs:
                    parts.append("关键文献：" + "；".join(refs[:10]))
                parts.append("国内现状：" + str(lit_review.get("domestic", "")))
                parts.append("国际现状：" + str(lit_review.get("international", "")))
                literature_context = "\n".join(parts)

        task = generate_proposal_task.delay(
            project_id=str(design.project_id) if design.project_id else "",
            design_id=payload.design_id,
            project_design=design.content or {},
            research_direction=direction,
            literature_context=literature_context,
        )
        return {"task_id": task.id, "status": "pending"}

    # --- 同步模式 ---
    # 加载项目设计
    design = db.query(ProjectDesign).filter(ProjectDesign.id == payload.design_id).first()
    if not design:
        raise HTTPException(status_code=404, detail="项目设计不存在")

    # 加载研究方向
    direction = None
    if design.direction_id:
        rd = db.query(ResearchDirection).filter(ResearchDirection.id == design.direction_id).first()
        if rd:
            direction = {
                "title": rd.title,
                "background": rd.background,
                "research_questions": rd.research_questions,
                "methods": rd.methods,
                "expected_outputs": rd.expected_outputs,
                "innovation": rd.innovation,
                "content": rd.content,
            }

    # 构建文献上下文
    literature_context = ""
    if isinstance(design.content, dict):
        lit_review = design.content.get("literature_review", {})
        if isinstance(lit_review, dict):
            parts = []
            refs = lit_review.get("key_references", [])
            if refs:
                parts.append("关键文献：" + "；".join(refs[:10]))
            parts.append("国内现状：" + str(lit_review.get("domestic", "")))
            parts.append("国际现状：" + str(lit_review.get("international", "")))
            literature_context = "\n".join(parts)

    # 调用 Agent 生成
    result = proposal_agent.generate(
        project_design=design.content or {},
        research_direction=direction,
        literature_context=literature_context,
        allowed_references=collect_allowed_references_from_design(design.content or {}),
    )

    # 保存到数据库
    proposal = Proposal(
        project_id=design.project_id,
        design_id=design.id,
        title=result.get("title", f"{design.topic} —— 开题报告"),
        content=result.get("sections", {}),
    )
    db.add(proposal)
    db.commit()
    db.refresh(proposal)

    # 生成 docx 并保存（MinIO + 本地双写）
    _ensure_storage_dir()
    filename = f"proposal_{proposal.id}.docx"
    filepath = os.path.join(STORAGE_DIR, filename)
    object_key = f"proposals/{filename}"
    try:
        buf = _build_docx(proposal)
        docx_bytes = buf.read()
        # 本地保存
        with open(filepath, "wb") as f:
            f.write(docx_bytes)
        # MinIO 保存
        try:
            from ..services.upload_service import save_bytes
            save_bytes(docx_bytes, object_key,
                       content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception:
            logger.warning("MinIO 保存提案失败，仅保存到本地")
        proposal.docx_path = filepath
        db.commit()
        db.refresh(proposal)
    except Exception as e:
        logger.warning(f"生成 docx 失败: {e}")

    return _proposal_to_out(proposal)


@router.get("/{proposal_id}", response_model=ProposalOut)
def get_proposal(proposal_id: str, db: Session = Depends(get_db)):
    """获取开题报告详情"""
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="开题报告不存在")
    return _proposal_to_out(proposal)


@router.get("/{proposal_id}/download")
def download_proposal(
    proposal_id: str,
    db: Session = Depends(get_db),
    format: str = Query("docx", pattern="^(docx|pdf)$"),
):
    """下载开题报告 docx 或 pdf 文件（MinIO 优先，本地 fallback）"""
    proposal = db.query(Proposal).filter(Proposal.id == proposal_id).first()
    if not proposal:
        raise HTTPException(status_code=404, detail="开题报告不存在")

    if format == "pdf":
        # PDF 实时生成
        filename = _safe_filename(proposal.title, f"proposal_{proposal_id}", "pdf")
        buf = io.BytesIO(_build_pdf(proposal))
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    # ---- DOCX（原有逻辑不变） ----
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    filename = _safe_filename(proposal.title, f"proposal_{proposal_id}", "docx")

    # MinIO 优先
    if proposal.docx_path:
        object_key = f"proposals/{os.path.basename(proposal.docx_path)}"
        try:
            from ..services.upload_service import get_object_stream
            minio_result = get_object_stream(object_key)
            if minio_result is not None:
                stream, size, ct = minio_result
                return StreamingResponse(
                    stream,
                    media_type=content_type,
                    headers={
                        "Content-Disposition": f'attachment; filename="{filename}"',
                        "Content-Length": str(size),
                    },
                )
        except Exception:
            pass

    # 本地 fallback
    if proposal.docx_path and os.path.exists(proposal.docx_path):
        return FileResponse(
            proposal.docx_path,
            media_type=content_type,
            filename=filename,
        )

    # 实时生成
    buf = _build_docx(proposal)
    return StreamingResponse(
        buf,
        media_type=content_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        },
    )
