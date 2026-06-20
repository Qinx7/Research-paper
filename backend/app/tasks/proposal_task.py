"""开题报告生成 Celery 异步任务"""
import io
import logging
import os

from ..core.celery_app import celery_app
from ..core.database import SessionLocal
from ..models.proposal import Proposal
from ..agents.proposal_agent import proposal_agent
from ..agents.workflows import run_generate_proposal_workflow

logger = logging.getLogger(__name__)

STORAGE_DIR = os.path.join(os.path.dirname(__file__), "..", "..", "storage", "proposals")


@celery_app.task(bind=True, max_retries=1, default_retry_delay=30)
def generate_proposal_task(
    self,
    project_id: str,
    design_id: str,
    project_design: dict,
    research_direction: dict | None = None,
    literature_context: str = "",
    user_id: str | None = None,
) -> dict:
    """异步生成开题报告，存入 DB 并生成 docx。

    返回可直接用于 ProposalOut 的序列化数据。
    """
    try:
        db = SessionLocal()
        try:
            workflow_result = run_generate_proposal_workflow(
                db=db,
                project_id=project_id,
                design_id=design_id,
                project_design=project_design,
                research_direction=research_direction,
                literature_context=literature_context,
                user_id=user_id,
                proposal_agent=proposal_agent,
                record_db=db,
            )
            proposal = workflow_result["proposal"]

            proposal_id = str(proposal.id)

            # 生成 docx（本地 + MinIO 双写）
            _ensure_storage_dir()
            buf = _build_docx(proposal)
            docx_bytes = buf.read()
            filepath = os.path.join(STORAGE_DIR, f"proposal_{proposal_id}.docx")
            with open(filepath, "wb") as f:
                f.write(docx_bytes)
            # MinIO 保存
            try:
                from ..services.upload_service import save_bytes
                save_bytes(docx_bytes, f"proposals/proposal_{proposal_id}.docx",
                           content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception:
                pass
            proposal.docx_path = filepath
            db.commit()

            return {
                "id": proposal_id,
                "title": proposal.title,
                "sections_count": workflow_result.get("sections_count", 0),
                "download_url": f"/api/proposal/{proposal_id}/download",
                "workflow_run_id": workflow_result.get("workflow_run_id"),
            }
        finally:
            db.close()

    except Exception as exc:
        if self.request.retries < self.max_retries:
            raise self.retry(exc=exc)
        raise


def _ensure_storage_dir():
    os.makedirs(STORAGE_DIR, exist_ok=True)


def _build_docx(proposal: Proposal) -> io.BytesIO:
    """将开题报告渲染为 docx 文件"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    from ..schemas.proposal import SECTION_KEYS, SECTION_LABELS

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 封面标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(60)
    title_para.paragraph_format.space_after = Pt(30)
    run = title_para.add_run(proposal.title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(40)
    run = sub_para.add_run("开题报告")
    run.font.size = Pt(16)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 章节正文
    sections_data = proposal.content or {}
    for key in SECTION_KEYS:
        sec = sections_data.get(key, {})
        sec_content = sec.get("content", "") if isinstance(sec, dict) else str(sec)

        heading = doc.add_heading(SECTION_LABELS.get(key, key), level=1)
        for run in heading.runs:
            run.font.name = "黑体"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(16)

        for para_text in sec_content.strip().split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(para_text)
            run.font.name = "宋体"
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
