"""上传资料文本解析与切分服务。"""
from __future__ import annotations

import io
import importlib.util
import os
import re
from dataclasses import dataclass, field

from docx import Document

from ..core.config import settings


class UnsupportedDocumentType(ValueError):
    """文件类型暂不支持解析。"""


class DocumentParseError(ValueError):
    """文件内容解析失败。"""


def _with_parse_stage(error: DocumentParseError, stage: str) -> DocumentParseError:
    setattr(error, "parse_stage", stage)
    return error


@dataclass(frozen=True)
class ParsedChunk:
    """解析后可入库的文本片段。"""

    index: int
    content: str
    meta: dict = field(default_factory=dict)


@dataclass(frozen=True)
class ParsedDocument:
    """解析后的文档正文和元数据。"""

    text: str
    source_type: str
    meta: dict


@dataclass(frozen=True)
class ParseAttempt:
    """单次解析策略的结果。"""

    parser: str
    text: str


SUPPORTED_PARSE_EXTENSIONS = {".txt", ".md", ".docx", ".pdf"}
OCR_MAX_PAGES = 20
OCR_DPI = 200


def extract_text_from_bytes(data: bytes, filename: str) -> ParsedDocument:
    """按文件扩展名从字节内容中提取正文。"""
    ext = os.path.splitext(filename or "")[1].lower()
    if ext not in SUPPORTED_PARSE_EXTENSIONS:
        raise UnsupportedDocumentType(f"不支持解析该文件类型：{ext or '未知类型'}")

    if ext in {".txt", ".md"}:
        text = _decode_text(data)
        cleaned = clean_extracted_text(text)
        document_kind = "general_text"
        return ParsedDocument(
            text=cleaned,
            source_type=ext,
            meta={
                "parser": "text",
                "document_kind": document_kind,
                "section_outline": detect_section_outline(cleaned, source_type=ext, document_kind=document_kind),
            },
        )

    if ext == ".docx":
        text = _extract_docx_text(data)
        cleaned = clean_extracted_text(text)
        outline = detect_section_outline(cleaned, source_type=ext, document_kind="structured_docx")
        document_kind = "structured_docx" if outline else "general_docx"
        return ParsedDocument(
            text=cleaned,
            source_type=ext,
            meta={
                "parser": "python-docx",
                "document_kind": document_kind,
                "section_outline": outline,
            },
        )

    return _extract_pdf_document(data)


def clean_extracted_text(text: str) -> str:
    """清洗解析文本，保留段落边界并移除多余空白。"""
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    lines = [line.strip() for line in normalized.split("\n")]
    return "\n".join(line for line in lines if line).strip()


def chunk_text(text: str, *, chunk_size: int = 1600, overlap: int = 200) -> list[ParsedChunk]:
    """按字符长度切分文本，保留少量重叠避免跨段信息断裂。"""
    cleaned = clean_extracted_text(text)
    if not cleaned:
        return []

    safe_size = max(1, chunk_size)
    safe_overlap = max(0, min(overlap, safe_size // 2))
    chunks: list[ParsedChunk] = []
    start = 0
    while start < len(cleaned):
        end = min(len(cleaned), start + safe_size)
        if end < len(cleaned):
            split_at = max(cleaned.rfind("\n", start, end), cleaned.rfind("。", start, end))
            if split_at > start + safe_size // 2:
                end = split_at + 1
        content = cleaned[start:end].strip()
        if content:
            chunks.append(ParsedChunk(index=len(chunks), content=content))
        if end >= len(cleaned):
            break
        start = max(0, end - safe_overlap)
    return chunks


def chunk_document(parsed: ParsedDocument, *, chunk_size: int = 1600, overlap: int = 200) -> list[ParsedChunk]:
    """优先按章节切分文档，再对单章节内容做长度切块。"""
    outline = parsed.meta.get("section_outline") or []
    if not outline:
        return chunk_text(parsed.text, chunk_size=chunk_size, overlap=overlap)

    chunks: list[ParsedChunk] = []
    for section in outline:
        section_content = clean_extracted_text(section.get("content", ""))
        if not section_content:
            continue
        inner_chunks = chunk_text(section_content, chunk_size=chunk_size, overlap=overlap)
        for chunk in inner_chunks:
            chunks.append(
                ParsedChunk(
                    index=len(chunks),
                    content=chunk.content,
                    meta={
                        "section_title": section.get("title"),
                        "section_level": section.get("level"),
                        "section_path": section.get("path") or [],
                    },
                )
            )
    return chunks or chunk_text(parsed.text, chunk_size=chunk_size, overlap=overlap)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("文本编码无法识别，请转换为 UTF-8 后重试。")


def _extract_docx_text(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError(f"DOCX 解析失败：{exc}"), "docx") from exc


def _extract_pdf_document(data: bytes) -> ParsedDocument:
    """按多策略顺序解析 PDF：先文本提取，再 OCR fallback。"""
    attempts: list[ParseAttempt] = []
    last_error: DocumentParseError | None = None
    for parser_name, parser in _pdf_parse_strategies():
        try:
            text = parser(data)
        except DocumentParseError as exc:
            last_error = exc
            text = ""
        attempts.append(ParseAttempt(parser=parser_name, text=text))
        cleaned = clean_extracted_text(text)
        if cleaned:
            document_kind = _detect_pdf_document_kind(cleaned)
            structured_content = _extract_structured_content(cleaned, document_kind)
            return ParsedDocument(
                text=cleaned,
                source_type=".pdf",
                meta={
                    "parser": _join_strategy_names([attempt.parser for attempt in attempts]),
                    "strategy_chain": [attempt.parser for attempt in attempts],
                    "document_kind": document_kind,
                    "structured_fields": _structured_fields_for_kind(document_kind),
                    "structured_content": structured_content["content"],
                    "structured_confidence": structured_content["confidence"],
                },
            )

    if last_error is not None:
        raise last_error
    raise _with_parse_stage(DocumentParseError("PDF 未提取到文字，OCR 识别结果为空。"), "ocr")


def _pdf_parse_strategies():
    strategies = [
        ("pypdf", _extract_pdf_text),
    ]
    if _has_pdfplumber_support():
        strategies.append(("pdfplumber", _extract_pdfplumber_text))
    strategies.append(("ocr", _ocr_pdf_text))
    return strategies


def _join_strategy_names(names: list[str]) -> str:
    return "+".join(names)


def _detect_pdf_document_kind(text: str) -> str:
    lowered = (text or "").lower()
    scholarly_markers = [
        "abstract",
        "keywords",
        "references",
        "introduction",
        "doi",
    ]
    matches = sum(1 for marker in scholarly_markers if marker in lowered)
    return "scholarly_pdf" if matches >= 2 else "general_pdf"


def detect_section_outline(text: str, *, source_type: str, document_kind: str) -> list[dict]:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    if not lines:
        return []

    outline: list[dict] = []
    current = None
    for line in lines:
        match = _match_section_heading(line, source_type=source_type, document_kind=document_kind)
        if match:
            if current is not None:
                current["content"] = clean_extracted_text("\n".join(current["content_lines"]))
                current.pop("content_lines", None)
                outline.append(current)
            level, title = match
            path = [title] if level == 1 or not outline else list(outline[-1].get("path", [])[: max(0, level - 1)]) + [title]
            current = {
                "title": title,
                "level": level,
                "path": path,
                "content_lines": [],
            }
            continue

        if current is not None:
            current["content_lines"].append(line)

    if current is not None:
        current["content"] = clean_extracted_text("\n".join(current["content_lines"]))
        current.pop("content_lines", None)
        outline.append(current)

    return [item for item in outline if item.get("title")]


def _match_section_heading(line: str, *, source_type: str, document_kind: str) -> tuple[int, str] | None:
    stripped = line.strip()
    if not stripped:
        return None

    if re.match(r"^第[一二三四五六七八九十百零〇\d]+章", stripped):
        return 1, stripped
    if re.match(r"^第[一二三四五六七八九十百零〇\d]+节", stripped):
        return 2, stripped
    if re.match(r"^\d+\.\d+\.\d+\s*", stripped):
        return 3, stripped
    if re.match(r"^\d+\.\d+\s*", stripped):
        return 2, stripped
    if re.match(r"^\d+\.\s*", stripped):
        return 1, stripped
    if re.match(r"^chapter\s+\d+", stripped, re.IGNORECASE):
        return 1, stripped
    if re.match(r"^\d+\s+[A-Z][A-Za-z].*", stripped):
        return 1, stripped

    if document_kind == "scholarly_pdf" or source_type == ".docx":
        lowered = stripped.lower()
        english_titles = {
            "abstract": 1,
            "keywords": 2,
            "introduction": 1,
            "related work": 1,
            "method": 1,
            "methods": 1,
            "experiment": 1,
            "experiments": 1,
            "results": 1,
            "conclusion": 1,
            "conclusions": 1,
            "references": 1,
        }
        if lowered in english_titles:
            return english_titles[lowered], stripped

    return None


def _structured_fields_for_kind(document_kind: str) -> list[str]:
    if document_kind == "scholarly_pdf":
        return ["title", "abstract", "sections", "references"]
    return []


def _extract_structured_content(text: str, document_kind: str) -> dict[str, dict]:
    if document_kind != "scholarly_pdf":
        return {"content": {}, "confidence": {}}

    title = _extract_title(text)
    abstract = _extract_abstract(text)
    references_text = _extract_references_text(text)

    return {
        "content": {
            "title": title,
            "abstract": abstract,
            "references_text": references_text,
            "references_list": _split_references_list(references_text),
        },
        "confidence": {
            "title": "medium" if title else "low",
            "abstract": "high" if abstract else "low",
            "references": "high" if references_text else "low",
        },
    }


def _extract_title(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    if not lines:
        return ""
    skip_prefixes = {"abstract", "摘要", "keywords", "关键词", "references", "参考文献", "doi"}
    for line in lines[:8]:
        lowered = line.lower()
        if any(lowered.startswith(prefix) for prefix in skip_prefixes):
            continue
        if len(line) < 8:
            continue
        return line
    return ""


def _extract_abstract(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    pattern = re.compile(
        r"(?:^|\n)(?:abstract|摘要)\s*\n?(.*?)(?=\n(?:keywords|关键词|introduction|1\.|references|参考文献)\b|$)",
        re.IGNORECASE | re.DOTALL,
    )
    matched = pattern.search(normalized)
    if not matched:
        return ""
    return clean_extracted_text(matched.group(1))


def _extract_references_text(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    pattern = re.compile(r"(?:^|\n)(?:references|reference|参考文献)\s*\n?(.*)$", re.IGNORECASE | re.DOTALL)
    matched = pattern.search(normalized)
    if not matched:
        return ""
    return clean_extracted_text(matched.group(1))


def _split_references_list(references_text: str) -> list[str]:
    cleaned = clean_extracted_text(references_text)
    if not cleaned:
        return []

    numbered = re.split(r"\n(?=\[\d+\]|\d+\.\s)", cleaned)
    items = [item.strip() for item in numbered if item.strip()]
    if len(items) > 1:
        return items[:20]

    lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
    return lines[:20]


def _has_pdfplumber_support() -> bool:
    return importlib.util.find_spec("pdfplumber") is not None


def _extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError("当前环境缺少 PDF 解析库 pypdf，请安装依赖后重试。"), "pypdf") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError(f"PDF 解析失败：{exc}"), "pypdf") from exc


def _extract_pdfplumber_text(data: bytes) -> str:
    try:
        import pdfplumber
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError("当前环境缺少 pdfplumber，无法执行版面补充解析。"), "pdfplumber") from exc

    try:
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        return "\n".join(pages).strip()
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError(f"pdfplumber 解析失败：{exc}"), "pdfplumber") from exc


def _ocr_pdf_text(data: bytes, *, max_pages: int = OCR_MAX_PAGES, dpi: int = OCR_DPI) -> str:
    """将扫描版 PDF 页面渲染为图片并调用 Tesseract OCR。"""
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError("当前环境缺少 OCR 依赖 PyMuPDF 或 pytesseract，请安装后重试。"), "ocr") from exc

    if settings.TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    try:
        pytesseract.get_tesseract_version()
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError("当前环境缺少 OCR 引擎 Tesseract，请先安装 Tesseract OCR 和中文语言包。"), "ocr") from exc

    try:
        doc = fitz.open(stream=data, filetype="pdf")
        page_count = min(len(doc), max_pages)
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        texts: list[str] = []
        for index in range(page_count):
            page = doc.load_page(index)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            try:
                page_text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            except Exception:
                page_text = pytesseract.image_to_string(image, lang="eng")
            if page_text.strip():
                texts.append(page_text)
        return "\n".join(texts)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise _with_parse_stage(DocumentParseError(f"PDF OCR 识别失败：{exc}"), "ocr") from exc
