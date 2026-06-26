import unittest
from unittest.mock import patch


class ProjectDocumentParseServiceTests(unittest.TestCase):
    def test_plain_text_is_cleaned_and_chunked_with_overlap(self):
        from app.services.document_parse_service import clean_extracted_text, chunk_text

        text = "第一段内容。\n\n\n第二段内容包含研究方法和实验数据。\n第三段内容。"
        cleaned = clean_extracted_text(text)
        chunks = chunk_text(cleaned, chunk_size=18, overlap=6)

        self.assertNotIn("\n\n\n", cleaned)
        self.assertGreaterEqual(len(chunks), 2)
        self.assertEqual(chunks[0].index, 0)
        self.assertIn("研究方法", "".join(chunk.content for chunk in chunks))
        self.assertTrue(all(chunk.content.strip() for chunk in chunks))

    def test_unsupported_extension_raises_clear_error(self):
        from app.services.document_parse_service import UnsupportedDocumentType, extract_text_from_bytes

        with self.assertRaises(UnsupportedDocumentType) as ctx:
            extract_text_from_bytes(b"demo", "image.png")

        self.assertIn("不支持解析", str(ctx.exception))

    def test_text_pdf_can_be_extracted_when_pypdf_is_installed(self):
        from fpdf import FPDF

        from app.services.document_parse_service import extract_text_from_bytes

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(40, 10, text="AI writing support evidence")
        data = bytes(pdf.output())

        parsed = extract_text_from_bytes(data, "evidence.pdf")

        self.assertEqual(parsed.source_type, ".pdf")
        self.assertIn("AI writing support evidence", parsed.text)

    def test_docx_extract_detects_section_outline_for_structured_document(self):
        import io
        from docx import Document

        from app.services.document_parse_service import extract_text_from_bytes, chunk_document

        doc = Document()
        doc.add_paragraph("第一章 绪论")
        doc.add_paragraph("绪论第一段。")
        doc.add_paragraph("1.1 研究背景")
        doc.add_paragraph("背景第一段。")
        buffer = io.BytesIO()
        doc.save(buffer)
        data = buffer.getvalue()

        parsed = extract_text_from_bytes(data, "paper.docx")
        chunks = chunk_document(parsed, chunk_size=50, overlap=10)

        self.assertEqual(parsed.source_type, ".docx")
        self.assertEqual(parsed.meta["document_kind"], "structured_docx")
        self.assertEqual(parsed.meta["section_outline"][0]["title"], "第一章 绪论")
        self.assertEqual(parsed.meta["section_outline"][1]["title"], "1.1 研究背景")
        self.assertEqual(chunks[0].meta["section_title"], "第一章 绪论")
        self.assertEqual(chunks[1].meta["section_title"], "1.1 研究背景")

    def test_scanned_pdf_falls_back_to_ocr_when_pypdf_extracts_no_text(self):
        from fpdf import FPDF

        from app.services.document_parse_service import extract_text_from_bytes

        pdf = FPDF()
        pdf.add_page()
        data = bytes(pdf.output())

        with patch(
            "app.services.document_parse_service._ocr_pdf_text",
            return_value="OCR recognized scanned content",
            create=True,
        ) as ocr:
            parsed = extract_text_from_bytes(data, "scan.pdf")

        self.assertEqual(parsed.source_type, ".pdf")
        self.assertEqual(parsed.meta["parser"], "pypdf+pdfplumber+ocr")
        self.assertIn("OCR recognized scanned content", parsed.text)
        ocr.assert_called_once()

    def test_scanned_pdf_reports_clear_error_when_ocr_engine_is_missing(self):
        from fpdf import FPDF

        from app.services.document_parse_service import DocumentParseError, extract_text_from_bytes

        pdf = FPDF()
        pdf.add_page()
        data = bytes(pdf.output())

        with patch(
            "app.services.document_parse_service._ocr_pdf_text",
            side_effect=DocumentParseError("当前环境缺少 OCR 引擎 Tesseract。"),
            create=True,
        ):
            with self.assertRaises(DocumentParseError) as ctx:
                extract_text_from_bytes(data, "scan.pdf")

        self.assertIn("OCR 引擎", str(ctx.exception))

    def test_pdf_parser_meta_contains_strategy_chain(self):
        from fpdf import FPDF

        from app.services.document_parse_service import extract_text_from_bytes

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(40, 10, text="Structured PDF content")
        data = bytes(pdf.output())

        parsed = extract_text_from_bytes(data, "article.pdf")

        self.assertEqual(parsed.meta["parser"], "pypdf")
        self.assertEqual(parsed.meta["strategy_chain"], ["pypdf"])

    def test_pdf_extract_uses_text_strategy_before_ocr_strategy(self):
        from app.services.document_parse_service import _extract_pdf_document

        calls = []

        with patch(
            "app.services.document_parse_service._extract_pdf_text",
            side_effect=lambda data: calls.append("text") or "",
            create=True,
        ), patch(
            "app.services.document_parse_service._ocr_pdf_text",
            side_effect=lambda data: calls.append("ocr") or "OCR content",
            create=True,
        ):
            parsed = _extract_pdf_document(b"pdf-bytes")

        self.assertEqual(calls, ["text", "ocr"])
        self.assertEqual(parsed.meta["strategy_chain"], ["pypdf", "pdfplumber", "ocr"])

    def test_pdf_extract_uses_pdfplumber_before_ocr_when_available(self):
        from app.services.document_parse_service import _extract_pdf_document

        calls = []

        with patch(
            "app.services.document_parse_service._extract_pdf_text",
            side_effect=lambda data: calls.append("text") or "",
            create=True,
        ), patch(
            "app.services.document_parse_service._has_pdfplumber_support",
            return_value=True,
            create=True,
        ), patch(
            "app.services.document_parse_service._extract_pdfplumber_text",
            side_effect=lambda data: calls.append("pdfplumber") or "pdfplumber content",
            create=True,
        ), patch(
            "app.services.document_parse_service._ocr_pdf_text",
            side_effect=lambda data: calls.append("ocr") or "OCR content",
            create=True,
        ):
            parsed = _extract_pdf_document(b"pdf-bytes")

        self.assertEqual(calls, ["text", "pdfplumber"])
        self.assertEqual(parsed.meta["strategy_chain"], ["pypdf", "pdfplumber"])

    def test_scholarly_pdf_meta_is_reserved_for_future_structured_parsing(self):
        from fpdf import FPDF

        from app.services.document_parse_service import extract_text_from_bytes

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, text="Abstract\nThis paper proposes a retrieval augmented generation method.\nReferences\n[1] Example")
        data = bytes(pdf.output())

        parsed = extract_text_from_bytes(data, "paper.pdf")

        self.assertEqual(parsed.meta["document_kind"], "scholarly_pdf")
        self.assertIn("title", parsed.meta["structured_fields"])
        self.assertIn("abstract", parsed.meta["structured_fields"])
        self.assertIn("references", parsed.meta["structured_fields"])
        self.assertEqual(
            parsed.meta["structured_content"]["title"],
            "This paper proposes a retrieval augmented generation method.",
        )
        self.assertEqual(
            parsed.meta["structured_content"]["abstract"],
            "This paper proposes a retrieval augmented generation method.",
        )
        self.assertEqual(
            parsed.meta["structured_content"]["references_text"],
            "[1] Example",
        )
        self.assertEqual(
            parsed.meta["structured_content"]["references_list"],
            ["[1] Example"],
        )
        self.assertEqual(parsed.meta["structured_confidence"]["title"], "medium")
        self.assertEqual(parsed.meta["structured_confidence"]["abstract"], "high")
        self.assertEqual(parsed.meta["structured_confidence"]["references"], "high")

    def test_detect_section_outline_recognizes_chinese_numeric_and_reference_sections(self):
        from app.services.document_parse_service import detect_section_outline

        text = "\n".join([
            "第一章 绪论",
            "这里是绪论内容。",
            "1.1 研究背景",
            "这里是背景内容。",
            "2 Method",
            "This is method section.",
            "References",
            "[1] Example reference",
        ])

        sections = detect_section_outline(text, source_type=".pdf", document_kind="scholarly_pdf")

        self.assertEqual(sections[0]["title"], "第一章 绪论")
        self.assertEqual(sections[0]["level"], 1)
        self.assertEqual(sections[1]["title"], "1.1 研究背景")
        self.assertEqual(sections[1]["level"], 2)
        self.assertEqual(sections[2]["title"], "2 Method")
        self.assertEqual(sections[3]["title"], "References")

    def test_chunk_document_uses_section_meta_when_outline_exists(self):
        from app.services.document_parse_service import ParsedDocument, chunk_document

        parsed = ParsedDocument(
            text="\n".join([
                "第一章 绪论",
                "绪论第一段。",
                "1.1 研究背景",
                "背景第一段。",
            ]),
            source_type=".pdf",
            meta={
                "document_kind": "scholarly_pdf",
                "section_outline": [
                    {
                        "title": "第一章 绪论",
                        "level": 1,
                        "path": ["第一章 绪论"],
                        "content": "绪论第一段。",
                    },
                    {
                        "title": "1.1 研究背景",
                        "level": 2,
                        "path": ["第一章 绪论", "1.1 研究背景"],
                        "content": "背景第一段。",
                    },
                ],
            },
        )

        chunks = chunk_document(parsed, chunk_size=50, overlap=10)

        self.assertEqual(chunks[0].meta["section_title"], "第一章 绪论")
        self.assertEqual(chunks[0].meta["section_level"], 1)
        self.assertEqual(chunks[1].meta["section_title"], "1.1 研究背景")
        self.assertEqual(chunks[1].meta["section_path"], ["第一章 绪论", "1.1 研究背景"])


if __name__ == "__main__":
    unittest.main()
